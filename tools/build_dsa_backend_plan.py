from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


OUT_DOCX = "Java_DSA_Backend_4_Month_Study_Plan.docx"
OUT_PDF = "Java_DSA_Backend_4_Month_Study_Plan.pdf"


duration_summary = [
    ["Source", "Videos / scope", "Visible duration", "Notes"],
    [
        "Kunal Kushwaha",
        "69 videos",
        "79h 35m 37s",
        "Java + DSA + interview preparation course; repo includes code samples, assignments, and notes.",
    ],
    [
        "Pepcoding raw visible total",
        "All pasted Pepcoding material",
        "84h 47m 49s",
        "Raw total counts repeated pasted entries.",
    ],
    [
        "Pepcoding cleaned total",
        "Duplicate Binary Tree repeats removed",
        "81h 59m 41s",
        "More useful planning total. Some first videos had no visible duration, so real total may be slightly higher.",
    ],
]

pepcoding_durations = [
    ["Pepcoding playlist", "Visible videos counted", "Duration"],
    ["Java Foundation", "74", "6h 7m 29s"],
    ["Functions & Arrays", "43", "3h 59m 45s"],
    ["2D Arrays", "19", "2h 16m 36s"],
    ["Recursion & Backtracking Level 1", "53", "8h 37m 27s"],
    ["Time & Space", "14 visible durations", "3h 21m 9s"],
    ["Linked List", "53 visible durations", "3h 19m 4s"],
    ["Stack & Queue", "56", "6h 58m 4s"],
    ["Generic Tree", "52 visible durations", "5h 32m 43s"],
    ["Binary Tree", "raw 61 / cleaned 43", "10h 43m raw / 7h 55m cleaned"],
    ["BST", "17 visible durations", "1h 28m 40s"],
    ["HashMap & Heap", "22 visible durations", "3h 39m 34s"],
    ["Dynamic Programming", "35 visible durations", "11h 54m 29s"],
    ["Graphs", "17 visible durations", "5h 53m 39s"],
    ["Backtracking Level 2", "38 visible durations", "10h 55m 35s"],
]

structured_order = [
    "Java Foundation",
    "Functions + Arrays",
    "2D Arrays",
    "Searching + Binary Search",
    "Sorting + Time Complexity",
    "Strings + StringBuilder + ArrayList",
    "Basic Maths + Bit Manipulation",
    "Recursion Level 1",
    "Backtracking basics",
    "OOP + Core Java",
    "Linked List",
    "Stack + Queue",
    "HashMap",
    "Heap / PriorityQueue",
    "Generic Tree",
    "Binary Tree",
    "Binary Search Tree",
    "Dynamic Programming",
    "Graphs",
    "Backtracking Level 2 selected problems",
    "Interview revision + mocks",
]

phases = [
    {
        "title": "Phase 1: Java Foundation",
        "source": "Pepcoding Java Foundation playlist",
        "context": "Use Pepcoding as the main source for basics, loops, input, number problems, GCD/LCM, prime, Fibonacci, and patterns.",
        "keep": [
            "variables",
            "conditionals",
            "loops",
            "input",
            "prime number",
            "Fibonacci",
            "count digits",
            "digits of number",
            "reverse number",
            "inverse number",
            "rotate number",
            "GCD/LCM",
            "prime factorization",
        ],
        "skip": [
            "too many pattern questions",
            "Print Z",
            "repeated beginner syntax videos if already clear",
            "do only 5-6 pattern questions, not all 20",
        ],
        "add": [
            "Kunal 5: Introduction to Programming, only if weak",
            "Kunal 6: Flowcharts & Pseudocode, only if weak",
            "Kunal 8: First Java Program, I/O, Debugging, Datatypes, only if weak",
            "Kunal 9: Conditionals and Loops, only if weak",
            "Kunal 11: Functions/Methods, only if weak",
        ],
        "must": ["Do not spend too much time here."],
    },
    {
        "title": "Phase 2: Functions + Arrays",
        "source": "Pepcoding Functions and Arrays Level 1",
        "context": "Use Pepcoding for foundation-heavy array clarity, then add Kunal for interview patterns.",
        "keep": [
            "functions",
            "number system basics",
            "arrays introduction",
            "memory",
            "span of array",
            "find element",
            "sum of two arrays",
            "difference of two arrays",
            "reverse array",
            "rotate array",
            "inverse array",
            "subarrays",
            "subsets",
            "binary search",
            "ceil/floor",
            "first/last index",
        ],
        "skip": [
            "base addition/subtraction/multiplication if number systems are already clear",
            "bar chart, only if you want loop practice",
        ],
        "add": [
            "Kunal 12: Arrays and ArrayList",
            "Kunal 13: Linear Search",
            "Kunal 14: Binary Search",
            "Kunal 15: Binary Search Interview Questions",
            "Kunal 16: Binary Search in 2D Arrays",
        ],
        "must": [
            "two pointers",
            "sliding window basics",
            "prefix sum",
            "Kadane's algorithm",
            "Dutch national flag",
            "binary search on answer",
            "rotated sorted array",
            "peak element",
        ],
    },
    {
        "title": "Phase 3: 2D Arrays",
        "source": "Pepcoding 2D Arrays Level 1",
        "context": "Build matrix traversal and implementation confidence without overinvesting in low-frequency problems.",
        "keep": [
            "matrix multiplication",
            "wave traversal",
            "spiral traversal",
            "exit point",
            "rotate 90 degree",
            "shell rotate",
            "diagonal traversal",
            "saddle point",
            "search in sorted 2D array",
        ],
        "skip": ["do not spend too much time on shell rotate; understand once and move on"],
        "add": ["Kunal 16: Binary Search in 2D Arrays"],
        "must": ["This add-on is enough for this phase."],
    },
    {
        "title": "Phase 4: Time & Space + Sorting",
        "source": "Pepcoding Time and Space Level 1",
        "context": "Use Pepcoding for sorting algorithms and complexity intuition, then add Cycle Sort for interview patterns.",
        "keep": [
            "bubble sort",
            "selection sort",
            "insertion sort",
            "merge two sorted arrays",
            "merge sort",
            "partitioning",
            "quick sort",
            "quick select",
            "count sort",
            "radix sort",
            "sort 01",
            "sort 012",
            "target sum pair",
            "pivot of sorted rotated array",
        ],
        "skip": [],
        "add": [
            "Kunal 20: Cycle Sort",
            "Kunal 24: Time and Space Complexity",
            "Kunal 59: Count Sort",
            "Kunal 60: Radix Sort",
        ],
        "must": [
            "Cycle Sort helps with missing number, duplicate number, set mismatch, and first missing positive.",
        ],
    },
    {
        "title": "Phase 5: Strings, StringBuilder, Maths, Bit Manipulation",
        "source": "Kunal as the strong add-on source",
        "context": "Pepcoding Level 1 can feel incomplete here for modern interviews, so make Kunal the main source for this phase.",
        "keep": [],
        "skip": [],
        "add": [
            "Kunal 21: Strings and StringBuilder",
            "Kunal 25: Bitwise Operators + Number Systems",
            "Kunal 26: Maths for DSA",
            "Kunal 52: StringBuffer",
            "Kunal 53: BigInteger & BigDecimal",
        ],
        "must": [
            "String immutability",
            "StringBuilder",
            "palindrome",
            "anagram",
            "frequency count",
            "substring logic",
            "XOR",
            "power of two",
            "find unique element",
            "GCD/LCM",
            "sieve",
            "modular arithmetic basics",
            "This phase is important for Java backend and coding rounds.",
        ],
    },
    {
        "title": "Phase 6: Recursion Level 1",
        "source": "Pepcoding Recursion and Backtracking Level 1",
        "context": "Pepcoding is the main source for recursion basics, recursion in arrays, subsequences, keypad, paths, permutations, encodings, flood fill, target sum, N-Queens, and Knight's Tour.",
        "keep": [
            "print increasing/decreasing",
            "factorial",
            "power linear/log",
            "print zig-zag",
            "Tower of Hanoi",
            "recursion in arrays",
            "first/last/all indices",
            "get subsequence",
            "keypad",
            "stair paths",
            "maze paths",
            "print subsequence",
            "print permutations",
            "print encodings",
            "flood fill",
            "target sum subsets",
            "N-Queens",
        ],
        "skip": [
            "Knight's Tour is optional for now",
            "too many maze variations; do enough to understand the pattern",
        ],
        "add": [
            "Kunal 23: Intro to Recursion",
            "Kunal 27: Recursion Level 1 Questions",
            "Kunal 28: Recursion Array Questions",
            "Kunal 32: Subset/Subsequence/String Questions",
            "Kunal 33: Permutations",
            "Kunal 35: Backtracking Maze",
            "Kunal 36: N-Queens, N-Knights, Sudoku",
        ],
        "must": ["Use Kunal mainly for interview mapping, not as a full second course."],
    },
    {
        "title": "Phase 7: OOP + Core Java",
        "source": "Kunal OOP series",
        "context": "This is missing as a proper deep module in Pepcoding and is compulsory for a Java backend shift.",
        "keep": [],
        "skip": [],
        "add": [
            "Kunal 37: OOP 1",
            "Kunal 38: OOP 2",
            "Kunal 39: OOP 3",
            "Kunal 40: OOP 4",
            "Kunal 41: OOP 5",
            "Kunal 42: OOP 6",
            "Kunal 43: OOP 7",
        ],
        "must": [
            "classes",
            "objects",
            "constructors",
            "static",
            "packages",
            "inheritance",
            "polymorphism",
            "encapsulation",
            "abstraction",
            "interfaces",
            "generics",
            "lambda",
            "exception handling",
            "collections",
            "Vector",
            "enums",
            "This phase is mandatory because the target is Java backend development, not only DSA problem solving.",
        ],
    },
    {
        "title": "Phase 8: Linked List",
        "source": "Pepcoding Linked List Level 1",
        "context": "Use Pepcoding for implementation and core operations, then add Kunal for interview bridging.",
        "keep": [
            "implementation",
            "add/remove/get",
            "reverse data iterative",
            "reverse pointer iterative",
            "kth from end",
            "middle",
            "merge two sorted lists",
            "merge sort linked list",
            "remove duplicates",
            "odd-even",
            "k reverse",
            "recursive reverse",
            "palindrome",
            "fold",
            "add two lists",
            "intersection point",
        ],
        "skip": [],
        "add": [
            "Kunal 44: Linked List Tutorial",
            "Kunal 45: Linked List Interview Questions",
        ],
        "must": [
            "cycle detection",
            "cycle starting node",
            "remove nth node",
            "reorder list",
            "add two numbers",
            "reverse in k-group, optional",
        ],
    },
    {
        "title": "Phase 9: Stack and Queue",
        "source": "Pepcoding Stack and Queue Level 1",
        "context": "Focus on bracket, monotonic stack, expression, interval, and adapter problems.",
        "keep": [
            "duplicate brackets",
            "balanced brackets",
            "next greater element",
            "stock span",
            "largest rectangle histogram",
            "sliding window maximum",
            "infix evaluation",
            "infix conversion",
            "celebrity problem",
            "postfix/prefix evaluation",
            "merge overlapping intervals",
            "min stack",
            "queue/stack adapters",
        ],
        "skip": [
            "too many adapter variations if implementation is already clear",
            "OOP mini swap examples because Kunal OOP is already included",
        ],
        "add": [
            "Kunal 46: Stacks and Queues Tutorial",
            "Kunal 47: Stacks and Queues Interview Questions",
        ],
        "must": [],
    },
    {
        "title": "Phase 10: HashMap and Heap",
        "source": "Pepcoding HashMap and Heaps Level 1",
        "context": "Do this before trees because it is highly useful in interviews.",
        "keep": [
            "hashmap intro",
            "highest frequency character",
            "common elements",
            "longest consecutive sequence",
            "heap intro",
            "k largest elements",
            "nearly sorted array",
            "median priority queue",
            "merge k sorted lists",
            "implement priority queue",
            "implement hashmap",
            "comparable vs comparator",
        ],
        "skip": [],
        "add": [
            "Kunal 56: Heap + PriorityQueue",
            "Kunal 57: HashMap & HashTable",
        ],
        "must": [
            "two sum",
            "group anagrams",
            "top k frequent elements",
            "subarray sum equals k",
            "longest consecutive sequence",
            "kth largest element",
            "merge k sorted lists",
        ],
    },
    {
        "title": "Phase 11: Generic Tree",
        "source": "Pepcoding Generic Trees Level 1",
        "context": "No major Kunal add-on is needed here.",
        "keep": [
            "construction",
            "display",
            "size/max/height",
            "traversals",
            "level order",
            "mirror",
            "remove leaves",
            "linearize",
            "find",
            "node to root path",
            "LCA",
            "distance between nodes",
            "similar/mirror/symmetric",
            "multisolver",
            "predecessor/successor",
            "ceil/floor",
            "kth largest",
            "diameter",
            "iterative preorder/postorder",
        ],
        "skip": [
            "Iterable and Iterator optional for now",
            "advanced generic-tree problems if time is low",
        ],
        "add": [],
        "must": [],
    },
    {
        "title": "Phase 12: Binary Tree",
        "source": "Pepcoding Binary Trees Level 1",
        "context": "The pasted Binary Tree list includes repeated entries. Use the cleaned useful content, roughly 7h 55m.",
        "keep": [
            "constructor",
            "display",
            "size/sum/max/height",
            "traversals",
            "level order",
            "iterative traversals",
            "node to root path",
            "print K levels down",
            "print nodes K distance away",
            "path to leaf",
            "remove leaves",
            "diameter",
            "tilt",
            "is BST",
            "balanced tree",
            "largest BST subtree",
            "construct tree from preorder/inorder",
            "construct from postorder/inorder",
            "bottom view",
        ],
        "skip": ["duplicate repeated entries in the pasted list"],
        "add": [
            "Kunal 49: Binary Trees + BST tutorial",
            "Kunal 55: Binary Tree Interview Questions",
            "Kunal 63: Binary Tree from Preorder & Inorder",
            "Kunal 64: Vertical Order Traversal",
            "Kunal 66: Two Sum IV",
            "Kunal 67: Kth Smallest in BST",
            "Kunal 68: Convert Binary Tree to DLL",
            "Kunal 69: Correct Swapped Nodes in BST",
        ],
        "must": [],
    },
    {
        "title": "Phase 13: BST",
        "source": "Pepcoding Binary Search Tree Level 1",
        "context": "Use Pepcoding for BST fundamentals, then add targeted Kunal interview problems.",
        "keep": [
            "constructor",
            "size/sum/max/min",
            "add node",
            "remove node",
            "replace with sum of larger",
            "LCA",
            "print in range",
            "target sum pair",
        ],
        "skip": [],
        "add": [
            "Kunal 67: Kth Smallest in BST",
            "Kunal 69: Correct Swapped Nodes in BST",
        ],
        "must": [],
    },
    {
        "title": "Phase 14: Dynamic Programming",
        "source": "Pepcoding Dynamic Programming Level 1",
        "context": "Kunal's 69-video list does not show a full dedicated DP module, so keep Pepcoding as the main source.",
        "keep": [
            "introduction to DP",
            "climbing stairs",
            "climbing stairs with jumps",
            "min moves",
            "min cost path",
            "goldmine",
            "target sum subset",
            "coin change combination",
            "coin change permutation",
            "0/1 knapsack",
            "unbounded knapsack",
            "count binary strings",
            "arrange buildings",
            "decode ways",
            "max sum non-adjacent",
            "paint house",
            "tiling",
            "friends pairing",
            "partition into subsets",
            "stock buy/sell variants",
        ],
        "skip": [
            "Highway Billboard optional",
            "very hard stock K transactions after basics",
        ],
        "add": [],
        "must": [],
    },
    {
        "title": "Phase 15: Graphs",
        "source": "Pepcoding Graphs Level 1",
        "context": "Pepcoding is stronger as the main source for graphs.",
        "keep": [
            "graph representation",
            "has path",
            "all paths",
            "DFS",
            "connected components",
            "is connected",
            "number of islands",
            "BFS",
            "cycle detection",
            "bipartite",
            "infection spread",
            "Dijkstra",
            "Prim's",
            "topological sort",
            "iterative DFS",
        ],
        "skip": [
            "Hamiltonian path/cycle optional for first pass",
            "Knight's Tour already optional from recursion",
        ],
        "add": ["Kunal 65: Word Ladder"],
        "must": [],
    },
    {
        "title": "Phase 16: Backtracking Level 2",
        "source": "Pepcoding Backtracking Level 2",
        "context": "This is not fully compulsory for the first job-prep pass. Do selected problems before applications.",
        "keep": [
            "abbreviations",
            "N-Queens branch and bound",
            "Sudoku",
            "palindrome partitioning",
            "word break",
            "remove invalid parentheses",
            "permutations",
            "combinations",
            "coin change recursion",
            "equal sum subset",
            "maximum number after K swaps, optional",
        ],
        "skip": [
            "crossword puzzle",
            "cryptarithmetic",
            "too many queen combination/permutation variations",
            "too many words K-selection variations",
            "Tug of War optional",
            "advanced partition K subsets optional",
        ],
        "add": [],
        "must": [
            "Backtracking Level 2 is almost 11 hours, so do not complete everything before job applications.",
        ],
    },
]

kunal_definitely_add = [
    ["Kunal lecture", "Why add"],
    ["12 Arrays and ArrayList", "Java ArrayList clarity"],
    ["14 Binary Search", "Better interview explanation"],
    ["15 Binary Search Interview Questions", "Very important"],
    ["20 Cycle Sort", "Missing interview pattern"],
    ["21 Strings and StringBuilder", "Important for Java"],
    ["24 Time and Space Complexity", "Strong Big-O clarity"],
    ["25 Bitwise + Number System", "Pepcoding not enough for interview"],
    ["26 Maths for DSA", "Useful for coding rounds"],
    ["37-43 OOP series", "Mandatory for Java backend"],
    ["45 Linked List Interview Questions", "Interview bridge"],
    ["47 Stack/Queue Interview Questions", "Interview bridge"],
    ["55 Binary Tree Interview Questions", "Very important"],
    ["56 Heap/PriorityQueue", "Java PriorityQueue clarity"],
    ["57 HashMap/HashTable", "Java HashMap clarity"],
    ["63-69 Tree/BST interview problems", "Good revision"],
]

kunal_skip = [
    ["Kunal lecture", "Reason"],
    ["1-3 intro/interview motivation", "Watch only if you want motivation"],
    ["4 Git/GitHub", "Important, but not DSA; do separately"],
    ["22 Pattern Questions", "Pepcoding already has patterns"],
    ["48 Tic Tac Toe", "Not needed for DSA/job prep"],
    ["50 AVL Tree", "Optional"],
    ["51 Segment Tree", "Optional for now"],
    ["54 File Handling", "Java useful, but not DSA priority"],
    ["58 Karp-Rabin", "Optional"],
    ["61 Huffman Coding", "Skip for fresher DSA"],
    ["62 Mo's Algorithm", "Skip for now"],
]

month_plan = [
    [
        "Month 1",
        "Java Foundation fast revision; Functions & Arrays; 2D Arrays; Searching; Sorting; Strings; Time complexity; Bit manipulation basics.",
        "Binary Search; Cycle Sort; Strings; Time Complexity; Bitwise.",
        "90-100 questions",
    ],
    [
        "Month 2",
        "Recursion Level 1; Backtracking basics; OOP from Kunal; Linked List.",
        "OOP 37-43; Recursion subset/permutation if needed; Linked List interview questions.",
        "180-200 total questions",
    ],
    [
        "Month 3",
        "Stack/Queue; HashMap; Heap; Generic Tree; Binary Tree.",
        "Stack/Queue interview questions; Heap/PriorityQueue; HashMap; Binary Tree interview questions.",
        "280-300 total questions",
    ],
    [
        "Month 4",
        "BST; DP Level 1 selected; Graphs Level 1 selected; Backtracking Level 2 selected; mocks and revision.",
        "BST/tree interview problems and selected graph problem Word Ladder.",
        "350-400 total questions + job applications",
    ],
]

final_format = [
    "Pepcoding Java Foundation",
    "Pepcoding Functions & Arrays + Kunal Binary Search + Cycle Sort",
    "Pepcoding 2D Arrays + Kunal Binary Search in 2D",
    "Pepcoding Time & Space + Kunal Time Complexity",
    "Kunal Strings + Bitwise + Maths",
    "Pepcoding Recursion Level 1 + Kunal Recursion selected",
    "Kunal OOP 1-7",
    "Pepcoding Linked List + Kunal Linked List Interview Questions",
    "Pepcoding Stack & Queue + Kunal Stack/Queue Interview Questions",
    "Pepcoding HashMap & Heap + Kunal HashMap + Heap",
    "Pepcoding Generic Tree",
    "Pepcoding Binary Tree + Kunal Binary Tree Interview Questions",
    "Pepcoding BST + Kunal BST interview questions",
    "Pepcoding DP Level 1 selected",
    "Pepcoding Graphs Level 1 selected",
    "Pepcoding Backtracking Level 2 selected",
    "LeetCode/GFG revision + mocks",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold


def docx_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_text(cell, value, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            if widths:
                cell.width = widths[c_idx]
    doc.add_paragraph()
    return table


def add_docx_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def create_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Java DSA + Backend 4-Month Study Plan")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Pepcoding as the main DSA source; Kunal Kushwaha as Java, OOP, and interview booster; LeetCode/GFG for final practice.")

    doc.add_heading("1. Source Strategy", level=1)
    add_docx_bullets(
        doc,
        [
            "Pepcoding = main DSA understanding source.",
            "Kunal Kushwaha = missing interview, OOP, and modern Java add-on source.",
            "LeetCode/GFG = final practice source.",
            "Do not randomly jump between playlists. Follow Pepcoding in controlled order and add selected Kunal lectures where Pepcoding is weak for fresher Java backend jobs.",
        ],
    )

    doc.add_heading("2. Total Video Hours Calculated", level=1)
    docx_table(doc, duration_summary, [Inches(1.5), Inches(1.6), Inches(1.3), Inches(2.1)])

    doc.add_heading("3. Pepcoding Playlist-Wise Duration", level=1)
    docx_table(doc, pepcoding_durations, [Inches(2.4), Inches(2.0), Inches(2.1)])
    doc.add_paragraph("Pepcoding's DSA Level 1 is publicly organized around Java Foundation, Functions/Arrays, Strings/StringBuilder/ArrayList, and DSA Level 1 modules.")

    doc.add_heading("4. Final Correct Structure", level=1)
    add_docx_bullets(doc, structured_order)

    doc.add_heading("5. Playlist-Wise Exact Plan", level=1)
    for phase in phases:
        doc.add_heading(phase["title"], level=2)
        doc.add_paragraph(f"Source: {phase['source']}")
        doc.add_paragraph(phase["context"])
        rows = [["Area", "What to do"]]
        for key, label in [("keep", "Keep / complete"), ("skip", "Remove / reduce / optional"), ("add", "Add from Kunal"), ("must", "Must-cover / notes")]:
            values = phase[key]
            if values:
                rows.append([label, "; ".join(values)])
        docx_table(doc, rows, [Inches(1.8), Inches(4.7)])

    doc.add_heading("6. Kunal Lectures You Should Definitely Add", level=1)
    docx_table(doc, kunal_definitely_add, [Inches(2.5), Inches(4.0)])

    doc.add_heading("7. Kunal Lectures You Can Skip for Now", level=1)
    docx_table(doc, kunal_skip, [Inches(2.5), Inches(4.0)])

    doc.add_heading("8. Four-Month Execution Plan", level=1)
    docx_table(doc, [["Month", "Complete", "Kunal add-ons", "Target"]] + month_plan, [Inches(0.9), Inches(2.4), Inches(2.2), Inches(1.0)])

    doc.add_heading("9. Final Follow Format", level=1)
    for idx, item in enumerate(final_format, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
    doc.add_paragraph("Main rule: Pepcoding remains the main learning source because it is easier to understand. Kunal is the gap-filler and interview booster, not the main course.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Java DSA + Backend Study Plan")
    doc.save(OUT_DOCX)


styles = getSampleStyleSheet()
styles.add(ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=30, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER, spaceAfter=16))
styles.add(ParagraphStyle("CoverSub", parent=styles["BodyText"], fontName="Helvetica", fontSize=11, leading=16, alignment=TA_CENTER, textColor=colors.HexColor("#333333"), spaceAfter=18))
styles.add(ParagraphStyle("H1x", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=8))
styles.add(ParagraphStyle("H2x", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=10, spaceAfter=6))
styles.add(ParagraphStyle("H3x", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=colors.HexColor("#1F4D78"), spaceBefore=6, spaceAfter=4))
styles.add(ParagraphStyle("Bodyx", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12.5, textColor=colors.black, spaceAfter=5))
styles.add(ParagraphStyle("Smallx", parent=styles["BodyText"], fontName="Helvetica", fontSize=8, leading=10.5, textColor=colors.HexColor("#333333"), spaceAfter=3))
styles.add(ParagraphStyle("TableCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2, leading=9.2, textColor=colors.black))
styles.add(ParagraphStyle("TableHead", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7.4, leading=9.4, textColor=colors.HexColor("#0B2545")))


def p(text, style="Bodyx"):
    text = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(text, styles[style])


def bullets(items):
    return ListFlowable(
        [ListItem(p(item, "Bodyx"), leftIndent=10) for item in items],
        bulletType="bullet",
        leftIndent=16,
        bulletFontSize=7,
    )


def numbered(items):
    return ListFlowable(
        [ListItem(p(item, "Bodyx"), leftIndent=10) for item in items],
        bulletType="1",
        leftIndent=18,
    )


def pdf_table(rows, col_widths=None, font_size=7.2):
    data = []
    for r, row in enumerate(rows):
        style = "TableHead" if r == 0 else "TableCell"
        data.append([p(cell, style) for cell in row])
    table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9AA8B6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def phase_block(phase):
    rows = [["Area", "What to do"]]
    for key, label in [
        ("keep", "Keep / complete"),
        ("skip", "Remove / reduce / optional"),
        ("add", "Add from Kunal"),
        ("must", "Must-cover / notes"),
    ]:
        if phase[key]:
            rows.append([label, "; ".join(phase[key])])
    return KeepTogether(
        [
            p(phase["title"], "H2x"),
            p(f"Source: {phase['source']}", "Smallx"),
            p(phase["context"], "Bodyx"),
            pdf_table(rows, [1.35 * inch, 5.15 * inch]),
            Spacer(1, 8),
        ]
    )


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#555555"))
    canvas.drawString(inch, 0.45 * inch, "Java DSA + Backend Study Plan")
    canvas.drawRightString(7.5 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def create_pdf():
    pdf = BaseDocTemplate(
        OUT_PDF,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
    )
    frame = Frame(pdf.leftMargin, pdf.bottomMargin, pdf.width, pdf.height, id="normal")
    pdf.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=footer)])
    story = []

    story.append(Spacer(1, 1.2 * inch))
    story.append(p("Java DSA + Backend 4-Month Study Plan", "CoverTitle"))
    story.append(p("Pepcoding as the main DSA understanding source. Kunal Kushwaha as the missing interview, OOP, and modern Java add-on source. LeetCode/GFG as the final practice source.", "CoverSub"))
    story.append(Spacer(1, 0.4 * inch))
    story.append(pdf_table([["Main Rule"], ["Do not randomly jump between playlists. Follow Pepcoding in a controlled order, and wherever Pepcoding is weak or missing for fresher Java backend jobs, add selected Kunal lectures."]], [6.5 * inch]))
    story.append(PageBreak())

    story.append(p("1. Source Strategy", "H1x"))
    story.append(bullets([
        "Pepcoding = main DSA understanding source.",
        "Kunal Kushwaha = missing interview/OOP/modern Java add-on source.",
        "LeetCode/GFG = final practice source.",
        "Kunal is not the main course now; Kunal is the gap-filler and interview booster.",
    ]))
    story.append(Spacer(1, 8))

    story.append(p("2. Total Video Hours Calculated", "H1x"))
    story.append(pdf_table(duration_summary, [1.25 * inch, 1.45 * inch, 1.1 * inch, 2.7 * inch]))
    story.append(Spacer(1, 8))

    story.append(p("3. Pepcoding Playlist-Wise Duration", "H1x"))
    story.append(pdf_table(pepcoding_durations, [2.55 * inch, 1.8 * inch, 2.15 * inch]))
    story.append(p("Pepcoding's DSA Level 1 playlist is publicly organized around Java Foundation, Functions/Arrays, Strings/StringBuilder/ArrayList, and DSA Level 1 modules.", "Smallx"))

    story.append(p("4. Final Correct Structure", "H1x"))
    story.append(numbered(structured_order))
    story.append(PageBreak())

    story.append(p("5. Playlist-Wise Exact Plan: Keep, Skip, Add From Kunal", "H1x"))
    for phase in phases:
        story.append(phase_block(phase))

    story.append(PageBreak())
    story.append(p("6. Kunal Lectures You Should Definitely Add", "H1x"))
    story.append(pdf_table(kunal_definitely_add, [2.45 * inch, 4.05 * inch]))
    story.append(Spacer(1, 8))

    story.append(p("7. Kunal Lectures You Can Skip for Now", "H1x"))
    story.append(pdf_table(kunal_skip, [2.45 * inch, 4.05 * inch]))
    story.append(Spacer(1, 8))

    story.append(p("8. Final 4-Month Execution Plan", "H1x"))
    story.append(pdf_table([["Month", "Complete", "Kunal add-ons", "Target"]] + month_plan, [0.75 * inch, 2.15 * inch, 2.35 * inch, 1.25 * inch]))
    story.append(Spacer(1, 8))

    story.append(p("9. Final Follow Format", "H1x"))
    for idx, item in enumerate(final_format, 1):
        story.append(p(f"{idx}. {item}", "Bodyx"))
    story.append(Spacer(1, 8))
    story.append(pdf_table([["Final Rule"], ["Pepcoding remains the main learning source because you understand it better. Kunal is the gap-filler and interview booster. Finish with LeetCode/GFG revision and mocks."]], [6.5 * inch]))

    pdf.build(story)


if __name__ == "__main__":
    create_docx()
    create_pdf()
    print(OUT_DOCX)
    print(OUT_PDF)
